from __future__ import annotations

import json
import math
import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documents" / "generated_diploma"
IMG_DIR = OUT_DIR / "images"
OUT_DOCX = ROOT / "documents" / "Диплом Абрамович Никита Сергеевич - пояснительная записка.docx"
PAGE_MAP_PATH = OUT_DIR / "page_map.json"


SECTION_TITLES = [
    "ПЕРЕЧЕНЬ УСЛОВНЫХ ОБОЗНАЧЕНИЙ",
    "ВВЕДЕНИЕ",
    "1 АНАЛИТИЧЕСКИЙ РАЗДЕЛ",
    "1.1 Описание предметной области",
    "1.2 Постановка задачи",
    "1.3 Обзор систем-аналогов",
    "1.4 Анализ требований к веб-приложению",
    "1.5 Обоснование выбора технологического стека",
    "1.5.1 Обоснование выбора архитектуры и сервера",
    "1.5.2 Обоснование выбора СУБД",
    "2 ПРАКТИЧЕСКИЙ РАЗДЕЛ",
    "2.1 Функциональная модель",
    "2.1.1 Диаграмма вариантов использования",
    "2.1.2 Диаграмма активности",
    "2.1.3 Диаграмма последовательности",
    "2.2 Проектирование структуры базы данных",
    "2.3 Архитектура программного решения",
    "2.4 Тестирование программного продукта",
    "3 РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ",
    "4 ЭКОНОМИЧЕСКИЙ РАЗДЕЛ",
    "5 РАЗДЕЛ ПО ОХРАНЕ ТРУДА",
    "ЗАКЛЮЧЕНИЕ",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    "ПРИЛОЖЕНИЕ А",
]


DEFAULT_TOC = {
    "ПЕРЕЧЕНЬ УСЛОВНЫХ ОБОЗНАЧЕНИЙ": 5,
    "ВВЕДЕНИЕ": 6,
    "1 АНАЛИТИЧЕСКИЙ РАЗДЕЛ": 9,
    "1.1 Описание предметной области": 9,
    "1.2 Постановка задачи": 12,
    "1.3 Обзор систем-аналогов": 14,
    "1.4 Анализ требований к веб-приложению": 17,
    "1.5 Обоснование выбора технологического стека": 20,
    "1.5.1 Обоснование выбора архитектуры и сервера": 22,
    "1.5.2 Обоснование выбора СУБД": 23,
    "2 ПРАКТИЧЕСКИЙ РАЗДЕЛ": 25,
    "2.1 Функциональная модель": 25,
    "2.1.1 Диаграмма вариантов использования": 26,
    "2.1.2 Диаграмма активности": 28,
    "2.1.3 Диаграмма последовательности": 30,
    "2.2 Проектирование структуры базы данных": 32,
    "2.3 Архитектура программного решения": 36,
    "2.4 Тестирование программного продукта": 40,
    "3 РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ": 43,
    "4 ЭКОНОМИЧЕСКИЙ РАЗДЕЛ": 51,
    "5 РАЗДЕЛ ПО ОХРАНЕ ТРУДА": 57,
    "ЗАКЛЮЧЕНИЕ": 61,
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ": 63,
    "ПРИЛОЖЕНИЕ А": 65,
}


SOURCES = [
    "Методические указания по выполнению дипломных проектов. Специальность 2-40 01 01 «Программное обеспечение информационных технологий». – Барановичи : УО «Барановичский технологический колледж» Белкоопсоюза, 2023.",
    "ГОСТ 2.105-2019. Единая система конструкторской документации. Общие требования к текстовым документам. – Москва : Стандартинформ, 2019.",
    "ГОСТ 7.32-2017. Система стандартов по информации, библиотечному и издательскому делу. Отчет о научно-исследовательской работе. Структура и правила оформления. – Москва : Стандартинформ, 2017.",
    "React documentation [Электронный ресурс]. – Режим доступа: https://react.dev/. – Дата доступа: 07.06.2026.",
    "Vite documentation [Электронный ресурс]. – Режим доступа: https://vite.dev/. – Дата доступа: 07.06.2026.",
    "TypeScript documentation [Электронный ресурс]. – Режим доступа: https://www.typescriptlang.org/docs/. – Дата доступа: 07.06.2026.",
    "FastAPI documentation [Электронный ресурс]. – Режим доступа: https://fastapi.tiangolo.com/. – Дата доступа: 07.06.2026.",
    "SQLAlchemy documentation [Электронный ресурс]. – Режим доступа: https://docs.sqlalchemy.org/. – Дата доступа: 07.06.2026.",
    "SQLite documentation [Электронный ресурс]. – Режим доступа: https://www.sqlite.org/docs.html. – Дата доступа: 07.06.2026.",
    "MDN Web Docs. Progressive web apps [Электронный ресурс]. – Режим доступа: https://developer.mozilla.org/. – Дата доступа: 07.06.2026.",
    "OWASP Foundation. Web Security Testing Guide [Электронный ресурс]. – Режим доступа: https://owasp.org/. – Дата доступа: 07.06.2026.",
    "Национальный правовой Интернет-портал Республики Беларусь [Электронный ресурс]. – Режим доступа: https://pravo.by/. – Дата доступа: 07.06.2026.",
    "Единый портал электронных услуг Республики Беларусь [Электронный ресурс]. – Режим доступа: https://portal.gov.by/. – Дата доступа: 07.06.2026.",
    "Министерство юстиции Республики Беларусь [Электронный ресурс]. – Режим доступа: https://minjust.gov.by/. – Дата доступа: 07.06.2026.",
    "Министерство по налогам и сборам Республики Беларусь [Электронный ресурс]. – Режим доступа: https://nalog.gov.by/. – Дата доступа: 07.06.2026.",
    "Министерство труда и социальной защиты Республики Беларусь [Электронный ресурс]. – Режим доступа: https://mintrud.gov.by/. – Дата доступа: 07.06.2026.",
    "Документация проекта «Белпомощник»: README.md, ARCHITECTURE.md, PROJECT_STATUS.md, API_CONTRACTS.md, TZ_IMPLEMENTATION_REPORT.md. – Локальный репозиторий проекта, 2026.",
    "Мартин, Р. Чистая архитектура. Искусство разработки программного обеспечения / Р. Мартин. – Санкт-Петербург : Питер, 2021. – 352 с.",
    "Зыков, С. В. Введение в теорию программирования: объектно-ориентированный подход / С. В. Зыков. – Москва : ИНТУИТ, 2020. – 200 с.",
    "Андруш, В. Г. Охрана труда : учебник / В. Г. Андруш, П. Т. Ткачёва, К. Д. Яшин. – Минск : РИПО, 2022. – 350 с.",
    "Пивоварчик, А. А. Охрана труда : учебно-методическое пособие / А. А. Пивоварчик. – Гродно : ГрГУ, 2022. – 440 с.",
]


def ref(n: int) -> str:
    return f"[{n}]"


def ensure_dirs():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)


def safe_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/Library/Fonts/Times New Roman Bold.ttf" if bold else "/Library/Fonts/Times New Roman.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]
    for c in candidates:
        if Path(c).exists():
            return ImageFont.truetype(c, size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def rounded(draw, xy, radius, fill, outline=(46, 78, 122), width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color=(55, 85, 130), width=3):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    for delta in (math.pi * 0.82, -math.pi * 0.82):
        x = end[0] + length * math.cos(angle + delta)
        y = end[1] + length * math.sin(angle + delta)
        draw.line([end, (x, y)], fill=color, width=width)


def draw_center_text(draw, box, text, font, fill=(17, 24, 39)):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 24)
    line_h = int(font.size * 1.25)
    total_h = line_h * len(lines)
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=font, fill=fill)
        y += line_h


def make_diagrams():
    title_font = safe_font(30, True)
    box_font = safe_font(22, True)
    small_font = safe_font(18)
    blue = (37, 99, 235)
    pale = (239, 246, 255)
    green = (22, 163, 74)

    # Use Case
    img = Image.new("RGB", (1500, 920), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Диаграмма вариантов использования платформы «Белпомощник»", font=title_font, fill=(17, 24, 39))
    actors = {
        "Гражданин": (120, 190),
        "Редактор контента": (120, 520),
        "Администратор": (1230, 360),
    }
    for name, (x, y) in actors.items():
        d.ellipse((x + 30, y, x + 90, y + 60), outline=blue, width=4)
        d.line((x + 60, y + 60, x + 60, y + 150), fill=blue, width=4)
        d.line((x + 10, y + 95, x + 110, y + 95), fill=blue, width=4)
        d.line((x + 60, y + 150, x + 20, y + 220), fill=blue, width=4)
        d.line((x + 60, y + 150, x + 100, y + 220), fill=blue, width=4)
        d.text((x - 25, y + 235), name, font=small_font, fill=(17, 24, 39))
    use_cases = [
        ("Искать проблему", 430, 145),
        ("Выбрать жизненный сценарий", 790, 145),
        ("Создать личную ситуацию", 610, 285),
        ("Отмечать задачи", 430, 435),
        ("Добавлять документы", 790, 435),
        ("Просматривать новости", 610, 585),
        ("Редактировать контент", 430, 730),
        ("Управлять ролями", 790, 730),
    ]
    for text, x, y in use_cases:
        d.ellipse((x, y, x + 260, y + 92), fill=pale, outline=blue, width=3)
        draw_center_text(d, (x, y, x + 260, y + 92), text, small_font)
    for _, x, y in use_cases[:6]:
        d.line((230, 300, x, y + 46), fill=(148, 163, 184), width=2)
    for _, x, y in use_cases[6:7]:
        d.line((230, 630, x, y + 46), fill=(148, 163, 184), width=2)
    d.line((1230, 510, 1050, 776), fill=(148, 163, 184), width=2)
    d.line((1230, 510, 690, 776), fill=(148, 163, 184), width=2)
    img.save(IMG_DIR / "use_case.png", quality=95)

    # Activity
    img = Image.new("RGB", (1500, 980), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Диаграмма активности создания личной ситуации", font=title_font, fill=(17, 24, 39))
    steps = [
        ("Открытие главной страницы", 545, 115),
        ("Поиск или выбор категории", 545, 235),
        ("Просмотр проблемы/сценария", 545, 355),
        ("Создание личной ситуации", 545, 475),
        ("Формирование задач и документов", 545, 595),
        ("Отметка выполнения задач", 545, 715),
        ("Автоматический пересчет прогресса", 545, 835),
    ]
    d.ellipse((708, 82, 792, 166), fill=blue)
    prev = (750, 166)
    for i, (text, x, y) in enumerate(steps):
        rounded(d, (x, y, x + 410, y + 72), 22, fill=(248, 250, 252), outline=blue)
        draw_center_text(d, (x, y, x + 410, y + 72), text, box_font)
        if i == 0:
            arrow(d, prev, (750, y), blue)
        elif i > 0:
            arrow(d, (750, steps[i - 1][2] + 72), (750, y), blue)
    d.ellipse((703, 920, 797, 960), fill=(17, 24, 39))
    arrow(d, (750, 907), (750, 920), blue)
    img.save(IMG_DIR / "activity.png", quality=95)

    # Sequence
    img = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Диаграмма последовательности: создание ситуации из шаблона", font=title_font, fill=(17, 24, 39))
    lanes = [("Пользователь", 150), ("React UI", 430), ("Store/API client", 720), ("FastAPI", 1010), ("SQLite", 1270)]
    for name, x in lanes:
        rounded(d, (x - 95, 110, x + 95, 160), 18, fill=pale, outline=blue)
        draw_center_text(d, (x - 95, 110, x + 95, 160), name, small_font)
        d.line((x, 160, x, 820), fill=(203, 213, 225), width=3)
    messages = [
        (150, 430, 210, "Открывает сценарий"),
        (430, 720, 285, "createSituation(templateId)"),
        (720, 1010, 360, "POST /api/user/situations"),
        (1010, 1270, 435, "Сохранение ситуации и задач"),
        (1270, 1010, 510, "id ситуации, задачи"),
        (1010, 720, 585, "DTO ситуации"),
        (720, 430, 660, "Обновление состояния"),
        (430, 150, 735, "Переход в «Мои ситуации»"),
    ]
    for x1, x2, y, text in messages:
        arrow(d, (x1, y), (x2, y), green if x2 < x1 else blue, width=3)
        d.text((min(x1, x2) + 18, y - 32), text, font=small_font, fill=(17, 24, 39))
    img.save(IMG_DIR / "sequence.png", quality=95)

    # Architecture
    img = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Архитектура программного решения", font=title_font, fill=(17, 24, 39))
    boxes = [
        ("Web / PWA", 110, 170, 390, 290),
        ("Mobile WebView", 110, 370, 390, 490),
        ("React/Vite frontend", 600, 230, 940, 390),
        ("API client и store", 600, 500, 940, 640),
        ("FastAPI backend", 1090, 250, 1390, 400),
        ("SQLite MVP", 1090, 520, 1390, 650),
    ]
    for text, x1, y1, x2, y2 in boxes:
        rounded(d, (x1, y1, x2, y2), 24, fill=(248, 250, 252), outline=blue)
        draw_center_text(d, (x1, y1, x2, y2), text, box_font)
    arrow(d, (390, 230), (600, 300), blue)
    arrow(d, (390, 430), (600, 300), blue)
    arrow(d, (770, 390), (770, 500), blue)
    arrow(d, (940, 570), (1090, 325), blue)
    arrow(d, (1240, 400), (1240, 520), blue)
    d.text((595, 735), "Mock fallback сохраняет работоспособность интерфейса при недоступном API", font=small_font, fill=(71, 85, 105))
    img.save(IMG_DIR / "architecture.png", quality=95)

    # Database
    img = Image.new("RGB", (1500, 950), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 35), "Схема основных сущностей базы данных", font=title_font, fill=(17, 24, 39))
    entities = [
        ("users", ["id", "email", "role", "profile_id"], 80, 130),
        ("user_profiles", ["region", "city", "addresses", "settings"], 80, 385),
        ("scenario_templates", ["id", "title", "category", "status"], 540, 130),
        ("scenario_stages", ["id", "scenario_id", "order"], 540, 385),
        ("scenario_tasks", ["id", "stage_id", "deadline", "documents"], 540, 640),
        ("user_situations", ["id", "user_id", "template_id", "progress"], 1000, 130),
        ("user_tasks", ["id", "situation_id", "task_id", "is_done"], 1000, 385),
        ("documents", ["id", "user_id", "type", "expires_at"], 1000, 640),
    ]
    for name, fields, x, y in entities:
        rounded(d, (x, y, x + 330, y + 165), 20, fill=(248, 250, 252), outline=blue)
        d.text((x + 20, y + 18), name, font=box_font, fill=blue)
        for i, field in enumerate(fields):
            d.text((x + 28, y + 58 + i * 25), field, font=small_font, fill=(17, 24, 39))
    links = [
        ((410, 212), (540, 212)),
        ((245, 295), (245, 385)),
        ((705, 295), (705, 385)),
        ((705, 550), (705, 640)),
        ((870, 212), (1000, 212)),
        ((1165, 295), (1165, 385)),
        ((1165, 550), (1165, 640)),
    ]
    for a, b in links:
        arrow(d, a, b, (71, 85, 105), 3)
    img.save(IMG_DIR / "database.png", quality=95)


def crop_screenshot(src: Path, dst: Path):
    if not src.exists():
        return
    img = Image.open(src).convert("RGB")
    w, h = img.size
    # Keep desktop pages readable in A4 by cropping top meaningful viewport.
    if w > 800:
        img = img.crop((0, 0, w, min(h, 820)))
    img.save(dst, quality=95)


def prepare_screenshots():
    mapping = [
        ("09-desktop-light.png", "screen_home.png"),
        ("03-desktop-light.png", "screen_catalog.png"),
        ("05-desktop-light.png", "screen_situations.png"),
        ("06-desktop-light.png", "screen_documents.png"),
        ("11-desktop-light.png", "screen_admin.png"),
    ]
    for src_name, dst_name in mapping:
        crop_screenshot(ROOT / "docs" / "redesign-progress" / src_name, IMG_DIR / dst_name)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 12, center: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(size)
    r.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm: list[float]):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(Cm(width).twips)))
                tc_w.set(qn("w:type"), "dxa")


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def set_page_start(section, start: int):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def setup_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = Pt(18)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    for name, size, bold, before, after, align, indent in [
        ("Diploma Section", 16, True, 0, 18, WD_ALIGN_PARAGRAPH.CENTER, 0),
        ("Diploma Heading 2", 14, True, 12, 8, WD_ALIGN_PARAGRAPH.LEFT, 1.25),
        ("Diploma Heading 3", 14, True, 8, 4, WD_ALIGN_PARAGRAPH.LEFT, 1.25),
        ("Diploma Caption", 12, False, 4, 8, WD_ALIGN_PARAGRAPH.CENTER, 0),
        ("Diploma Table Caption", 12, False, 8, 4, WD_ALIGN_PARAGRAPH.LEFT, 0),
        ("Diploma Source", 14, False, 0, 0, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.25),
        ("Diploma TOC", 14, False, 0, 0, WD_ALIGN_PARAGRAPH.LEFT, 0),
    ]:
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = bold
        sp = st.paragraph_format
        sp.alignment = align
        sp.first_line_indent = Cm(indent)
        sp.line_spacing = Pt(18)
        sp.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        sp.space_before = Pt(before)
        sp.space_after = Pt(after)


def new_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(2)
    sec.different_first_page_header_footer = True
    setup_styles(doc)
    return doc


def p(doc, text="", style=None, align=None, bold=False, italic=False, size=None, first_indent=True):
    par = doc.add_paragraph(style=style)
    if align is not None:
        par.alignment = align
    if not first_indent:
        par.paragraph_format.first_line_indent = Cm(0)
    run = par.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return par


def section_heading(doc, title):
    par = p(doc, title.upper(), "Diploma Section", first_indent=False)
    return par


def subheading(doc, title, level=2):
    return p(doc, title, "Diploma Heading 2" if level == 2 else "Diploma Heading 3")


def add_table_caption(doc, caption):
    return p(doc, caption, "Diploma Table Caption", first_indent=False)


def add_figure_caption(doc, caption):
    return p(doc, caption, "Diploma Caption", first_indent=False)


def add_doc_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths_cm)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=12, center=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAFD")
        set_cell_margins(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=12, center=False if len(str(value)) > 8 else True)
            set_cell_margins(cells[i])
    return table


def add_picture(doc, path: Path, width_cm=15.2):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.first_line_indent = Cm(0)
    run = par.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    return par


def page_break(doc):
    doc.add_page_break()


def title_page(doc):
    for text in [
        "Учреждение образования",
        "«Барановичский технологический колледж» Белкоопсоюза",
    ]:
        p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    for _ in range(3):
        p(doc, "")
    p(doc, "ДОПУЩЕН К ЗАЩИТЕ", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, first_indent=False)
    p(doc, "Заместитель директора", align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
    p(doc, "по учебной работе", align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
    p(doc, "___________ Е. К. Хамитова", align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
    p(doc, "«____» ____________ 2026 г.", align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
    for _ in range(3):
        p(doc, "")
    p(doc, "ДИПЛОМНЫЙ ПРОЕКТ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, first_indent=False)
    p(doc, "РАЗРАБОТКА ИНФОРМАЦИОННОЙ ВЕБ-СИСТЕМЫ ДЛЯ РЕШЕНИЯ АДМИНИСТРАТИВНО-БЫТОВЫХ ЗАДАЧ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, first_indent=False)
    p(doc, "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, first_indent=False)
    p(doc, "ДП 2-40 01 01 – 471.01.20.2.05 ПЗ", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
    for _ in range(9):
        p(doc, "")
    p(doc, "Разработал учащийся группы 471", align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False)
    p(doc, "Абрамович Н. С. __________________", align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False)
    p(doc, "Руководитель проекта __________________", align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False)
    for _ in range(3):
        p(doc, "")
    p(doc, "Барановичи, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)


def assignment_page(doc):
    page_break(doc)
    section_heading(doc, "ЗАДАНИЕ НА ДИПЛОМНЫЙ ПРОЕКТ")
    assignment = [
        "Обучающийся: Абрамович Никита Сергеевич. Курс IV. Учебная группа 471.",
        "Специальность 2-40 01 01 «Программное обеспечение информационных технологий».",
        "Тема дипломного проекта: Разработка информационной веб-системы для решения административно-бытовых задач.",
        "Срок окончания дипломного проекта: 18.06.2026.",
        "Исходные данные по дипломному проекту: материалы технического задания, текущая версия веб-платформы «Белпомощник», методические указания по выполнению дипломных проектов, рекомендации по экономическому разделу и материалы раздела охраны труда.",
        "Содержание пояснительной записки: перечень условных обозначений, введение, аналитический раздел, практический раздел, руководство пользователя, экономический раздел, раздел по охране труда, заключение, список использованных источников и приложение.",
        "Графическая часть: диаграмма вариантов использования, диаграмма активности, диаграмма последовательности, схема структуры базы данных, архитектурная схема программного решения, скриншоты основных страниц веб-системы.",
        "Консультант по экономическому разделу: С. М. Мазура.",
        "Консультант по охране труда: А. С. Бродюк.",
        "Консультант по соблюдению ТНПА: Н. А. Зинькович.",
    ]
    for text in assignment:
        p(doc, text)
    p(doc, "")
    p(doc, "Председатель цикловой комиссии _____________ Е. А. Герасимова", first_indent=False)
    p(doc, "Руководитель проекта ______________________", first_indent=False)
    p(doc, "Учащийся _________________________________", first_indent=False)


def abstract_page(doc):
    page_break(doc)
    section_heading(doc, "РЕФЕРАТ")
    for text in [
        "Пояснительная записка содержит описание разработки информационной веб-системы «Белпомощник», предназначенной для помощи гражданам Республики Беларусь при решении административно-бытовых задач. Работа включает аналитический раздел, практическое проектирование, руководство пользователя, экономическое обоснование и раздел по охране труда.",
        "Ключевые слова: ВЕБ-СИСТЕМА, БЕЛПОМОЩНИК, REACT, VITE, FASTAPI, SQLITE, ЖИЗНЕННЫЙ СЦЕНАРИЙ, ДОКУМЕНТЫ, УВЕДОМЛЕНИЯ, АДМИНИСТРАТИВНО-БЫТОВЫЕ ЗАДАЧИ, РОЛИ ПОЛЬЗОВАТЕЛЕЙ, WEBVIEW.",
        "Объектом разработки является веб-платформа «Белпомощник», которая предоставляет гражданину понятный маршрут решения бытовой или административной ситуации. Предметом разработки является программная реализация каталога проблем, жизненных сценариев, личных ситуаций, документов, уведомлений, профиля пользователя и административного наполнения контента.",
        "Цель дипломного проекта – разработать информационную веб-систему, которая помогает пользователю выбрать проблему или жизненный сценарий, получить пошаговый план действий, контролировать задачи и видеть связанные документы, сроки, учреждения и официальные источники.",
        "В ходе выполнения проекта подготовлена React/Vite-версия пользовательского интерфейса, backend-основа на FastAPI и SQLite, WebView-оболочка для мобильного показа, роли пользователей, гостевой режим, каталог проблем, жизненные сценарии, документы, уведомления, новости и административная панель.",
    ]:
        p(doc, text)


def toc_page(doc, page_map: dict[str, int]):
    page_break(doc)
    section_heading(doc, "СОДЕРЖАНИЕ")
    toc_items = [
        ("ПЕРЕЧЕНЬ УСЛОВНЫХ ОБОЗНАЧЕНИЙ", 1),
        ("ВВЕДЕНИЕ", 1),
        ("1 АНАЛИТИЧЕСКИЙ РАЗДЕЛ", 1),
        ("1.1 Описание предметной области", 2),
        ("1.2 Постановка задачи", 2),
        ("1.3 Обзор систем-аналогов", 2),
        ("1.4 Анализ требований к веб-приложению", 2),
        ("1.5 Обоснование выбора технологического стека", 2),
        ("1.5.1 Обоснование выбора архитектуры и сервера", 3),
        ("1.5.2 Обоснование выбора СУБД", 3),
        ("2 ПРАКТИЧЕСКИЙ РАЗДЕЛ", 1),
        ("2.1 Функциональная модель", 2),
        ("2.1.1 Диаграмма вариантов использования", 3),
        ("2.1.2 Диаграмма активности", 3),
        ("2.1.3 Диаграмма последовательности", 3),
        ("2.2 Проектирование структуры базы данных", 2),
        ("2.3 Архитектура программного решения", 2),
        ("2.4 Тестирование программного продукта", 2),
        ("3 РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ", 1),
        ("4 ЭКОНОМИЧЕСКИЙ РАЗДЕЛ", 1),
        ("5 РАЗДЕЛ ПО ОХРАНЕ ТРУДА", 1),
        ("ЗАКЛЮЧЕНИЕ", 1),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 1),
        ("ПРИЛОЖЕНИЕ А", 1),
    ]
    for title, level in toc_items:
        par = doc.add_paragraph(style="Diploma TOC")
        par.paragraph_format.first_line_indent = Cm(0)
        par.paragraph_format.left_indent = Cm(0.6 * (level - 1))
        run = par.add_run(title)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)
        page_run = par.add_run(" " + "." * max(3, 76 - len(title) - level * 3) + f" {page_map.get(title, DEFAULT_TOC.get(title, 0))}")
        page_run.font.name = "Times New Roman"
        page_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        page_run.font.size = Pt(14)


def abbreviations(doc):
    section_heading(doc, "ПЕРЕЧЕНЬ УСЛОВНЫХ ОБОЗНАЧЕНИЙ")
    pairs = [
        ("API", "Application Programming Interface, программный интерфейс приложения"),
        ("CRUD", "создание, чтение, изменение и удаление данных"),
        ("FastAPI", "web-фреймворк Python для разработки серверного API"),
        ("JWT", "токен авторизации, используемый при обмене данными между клиентом и сервером"),
        ("MVP", "минимально жизнеспособный продукт"),
        ("PWA", "progressive web application, прогрессивное веб-приложение"),
        ("RBAC", "управление доступом на основе ролей"),
        ("React", "библиотека JavaScript для построения пользовательских интерфейсов"),
        ("SQLite", "встроенная реляционная система управления базами данных"),
        ("UI", "пользовательский интерфейс"),
        ("UX", "пользовательский опыт"),
        ("WebView", "компонент для отображения веб-интерфейса внутри нативной оболочки"),
        ("БД", "база данных"),
        ("ЖКХ", "жилищно-коммунальное хозяйство"),
        ("ПК", "персональный компьютер"),
        ("ПЗ", "пояснительная записка"),
        ("СУБД", "система управления базами данных"),
        ("ТЗ", "техническое задание"),
    ]
    p(doc, "В настоящей пояснительной записке применяются следующие сокращения и условные обозначения:")
    for key, val in pairs:
        p(doc, f"{key} – {val}", first_indent=False)


def expand(doc, topic: str, points: list[str]):
    for point in points:
        p(doc, point)


EXTRA = {
    "domain": [
        "Административно-бытовые задачи отличаются тем, что пользователь чаще всего обращается к системе не из профессионального интереса, а в ситуации неопределённости. Он может не знать названия органа, правильного документа или юридического термина. Поэтому интерфейс должен начинаться не с ведомственной классификации, а с человеческой формулировки проблемы. В «Белпомощнике» это выражается через каталог понятных карточек и жизненных сценариев.",
        "При проектировании предметной области важно учитывать, что одна и та же ситуация может затрагивать несколько организаций. Например, рождение ребёнка связано с медицинской организацией, ЗАГСом, органами социальной защиты, регистрацией по месту жительства и поликлиникой. Если пользователь видит только один шаг, он не понимает всей цепочки. Поэтому сценарий хранится как набор этапов и задач, а не как один справочный текст.",
        "Система также должна помогать с подготовкой документов. В реальной жизни ошибка в одном документе может привести к повторному посещению учреждения, потере времени и стрессу. Поэтому документы в проекте связаны с задачами. Пользователь видит, какие документы нужны для конкретного действия, а не только общий список в конце статьи.",
        "Для административно-бытовой тематики важно разделять справочную и персональную информацию. Справочная часть объясняет общий порядок действий, а персональная часть относится к конкретному пользователю: его ситуации, документам, адресам, срокам и выполненным задачам. Такое разделение позволяет показывать публичную информацию гостю, но защищать личные действия.",
        "Ещё одной особенностью предметной области является изменчивость правил. Нормативные требования, формы заявлений и порядок обращения могут обновляться. Поэтому в проекте предусмотрены официальные источники и новости. Они не заменяют проверку на официальном ресурсе, но помогают пользователю увидеть, откуда взята информация и какие материалы требуют уточнения.",
        "Веб-платформа должна быть достаточно гибкой, чтобы поддерживать новые категории. Сегодня в системе представлены документы, семья, ЖКХ, налоги, работа и здоровье. В дальнейшем могут появиться образование, миграционные вопросы, социальная поддержка, транспорт и другие направления. Поэтому категории и сценарии рассматриваются как управляемые сущности, а не как жёстко зашитые страницы.",
        "Предметная область требует простого языка. Пользователь может быть подростком, пожилым человеком, родителем, предпринимателем или человеком, впервые столкнувшимся с административной процедурой. Поэтому формулировки должны объяснять действия без избыточной канцелярской нагрузки. Важная информация должна быть разбита на этапы и задачи.",
        "Наконец, система должна быть полезна не только конечному пользователю, но и ответственному за контент специалисту. Если новые материалы нельзя удобно добавлять и проверять, платформа быстро устареет. По этой причине в проект включена административная и редакторская часть, где можно управлять сценариями, публикациями, источниками и справочниками.",
    ],
    "task": [
        "Постановка задачи учитывает практический пользовательский путь. Сначала пользователь должен понять, что система вообще может помочь в его вопросе. Затем он должен найти подходящий материал, открыть его, оценить последовательность действий и при необходимости сохранить сценарий как личный план. Такой путь задаёт требования к навигации и структуре данных.",
        "Особое внимание уделяется гостевому режиму. Для открытой справочной платформы важно, чтобы пользователь мог познакомиться с системой без обязательной регистрации. Однако создание ситуации, добавление документов и изменение профиля уже являются персональными действиями. Поэтому система должна мягко объяснять необходимость входа, не блокируя просмотр публичных материалов.",
        "Постановка задачи также включает поддержку ролей. Гражданин использует систему для решения своей ситуации. Редактор контента наполняет материалы и источники. Администратор управляет пользователями и системными справочниками. Такое разделение необходимо, чтобы в дальнейшем платформа могла использоваться не только как учебный прототип, но и как управляемая информационная система.",
        "Для реализации личных ситуаций требуется преобразование шаблона в пользовательские данные. Шаблон хранит общий план, а пользовательская ситуация хранит состояние выполнения. Это позволяет нескольким пользователям использовать один сценарий независимо, не изменяя исходный контент.",
        "Важным элементом задачи является автоматический расчёт прогресса. Пользователь не должен вручную вводить процент выполнения. Прогресс определяется количеством выполненных задач по отношению к общему числу задач ситуации. Такой подход прост, понятен и проверяем.",
        "Система также должна поддерживать расширение через backend. На раннем этапе часть данных может храниться локально или в mock fallback, но архитектура должна позволять переносить данные на сервер. Поэтому API, схемы и модели проектируются заранее.",
        "Постановка задачи не включает предоставление юридической консультации. Система носит справочный характер и должна направлять пользователя к официальным источникам. Это снижает риск неправильного толкования информации и соответствует образовательному формату дипломного проекта.",
    ],
    "analogs": [
        "При сравнении с официальными порталами важно учитывать, что они обладают высокой достоверностью, но не всегда удобны для неподготовленного пользователя. Ведомственная структура логична для специалиста, однако гражданин чаще формулирует запрос иначе: «потерял паспорт», «родился ребёнок», «нужно открыть ИП». Поэтому «Белпомощник» использует пользовательскую формулировку как входную точку.",
        "Справочные сайты проще для восприятия, но их слабым местом часто является отсутствие персонального состояния. Пользователь прочитал статью и должен сам запомнить, что уже сделал. В «Белпомощнике» сценарий можно сохранить, а задачи отмечать выполненными. Это превращает информацию в инструмент сопровождения.",
        "Отдельные сервисы электронных услуг позволяют выполнить конкретную операцию, но не всегда объясняют весь путь вокруг неё. Например, подача заявления может быть только одним этапом более длинного сценария. Поэтому в проекте сценарий рассматривается шире, чем одна электронная форма.",
        "Планировщики задач удобны для контроля выполнения, но не знают предметной области. Они не подсказывают официальные источники, документы, учреждения и зависимости шагов. В дипломном проекте эти данные закладываются в шаблон сценария, что повышает прикладную ценность системы.",
        "Сравнение аналогов показывает, что главной особенностью «Белпомощника» является связь справочного контента с пользовательским прогрессом. Это не просто база статей и не просто список задач, а объединение двух подходов в одном интерфейсе.",
    ],
    "requirements": [
        "Функциональные требования проверяются через реальные сценарии. Если пользователь может открыть каталог, выбрать «Рождение ребёнка», создать личную ситуацию, увидеть задачи и отметить их выполненными, значит основная ценность системы подтверждается. Поэтому тестирование строится вокруг таких сквозных путей.",
        "Для профиля пользователя важны не только имя и электронная почта. Регион, город, район и адреса помогают персонализировать подсказки по учреждениям. Предпочтения источников позволяют поднять выше новости из выбранных официальных ресурсов. Эти функции готовят систему к более точной релевантности.",
        "Документы пользователя относятся к чувствительным данным. Даже в учебной версии необходимо демонстрировать аккуратный подход: маскирование реквизитов, предупреждение о сканах, разделение публичной и личной информации. В production-версии потребуется полноценное шифрование и политика хранения персональных данных.",
        "К интерфейсным требованиям относится работа на разных устройствах. Desktop-версия должна эффективно использовать ширину экрана, а мобильная версия – быть удобной для пальца. WebView-оболочка должна скрывать технические ошибки загрузки и показывать понятные сообщения о сети или сервере.",
        "К административным требованиям относится возможность наполнения системы без изменения кода. Редактор должен управлять содержанием, а администратор – системными параметрами. Это особенно важно для проекта, где справочная информация должна регулярно обновляться.",
        "Нефункциональные требования также включают сопровождаемость кода. Разделение frontend, backend, миграций, схем и сервисов облегчает развитие проекта. Если вся логика находится внутри страниц интерфейса, дальнейший переход к production становится сложнее.",
        "Требование к источникам связано с доверием. Пользователь должен видеть, что материал опирается на официальные ресурсы. Поэтому новости, сценарии и справочные карточки должны иметь связи с источниками, датами проверки и пояснениями.",
    ],
    "stack": [
        "Выбор React объясняется необходимостью построить насыщенный интерфейс с большим количеством состояний: фильтры, карточки, формы, модальные окна, роли, адаптивная навигация и защищённые действия. Компонентная модель позволяет переиспользовать элементы и удерживать визуальную целостность.",
        "Vite выбран как современный инструмент сборки, который ускоряет локальную разработку. Для дипломного проекта это практически важно: интерфейс часто менялся, проверялся в браузере и адаптировался под WebView. Быстрая сборка уменьшает время итерации.",
        "TypeScript помогает описать модели данных, используемые в интерфейсе. Для «Белпомощника» это особенно заметно в сценариях, задачах, документах и профиле. Ошибка в названии поля может привести к тому, что карточка не покажет важную информацию; типизация снижает такой риск.",
        "FastAPI выбран для backend-слоя из-за лаконичного описания API и удобной работы со схемами. Он подходит для учебного проекта, но при этом не является игрушечным решением: на его основе можно строить полноценные серверные сервисы.",
        "SQLite используется как стартовая СУБД. Для дипломного проекта это снижает сложность запуска и делает приложение воспроизводимым. Все основные данные могут быть созданы миграциями и seed-скриптами, что удобно при демонстрации.",
        "WebView-оболочка выбрана как промежуточный способ показать кроссплатформенность без поддержки отдельного нативного интерфейса. Это решение разумно, пока основной акцент делается на web-платформе и адаптивном интерфейсе.",
    ],
    "functional": [
        "Функциональная модель системы строится вокруг понятия действия. Пользователь не просто читает материал, а выполняет последовательность действий: ищет, выбирает, создаёт, отмечает, возвращается к прогрессу. Поэтому каждая страница должна вести к следующему осмысленному шагу.",
        "Для гостя модель должна быть достаточно открытой. Если скрыть слишком много, пользователь не поймёт ценность продукта. Если разрешить всё, возникнет риск смешения данных. Поэтому гость видит структуру сервиса, но персональные операции переводят его к входу.",
        "Редакторская роль нужна для жизнеспособности системы. Без неё новые сценарии и источники можно добавлять только разработчиком, что противоречит идее информационной платформы. В дипломной версии редакторская часть показывает концепцию наполнения.",
        "Администраторская роль отличается от редакторской. Администратор работает не только с контентом, но и с пользователями, ролями, справочниками и аналитикой. Такое разделение соответствует будущему развитию системы.",
        "Диаграммы в практическом разделе показывают не все возможные детали, а наиболее важные связи. Для дипломного проекта достаточно отразить центральный пользовательский сценарий и общую архитектуру, чтобы была понятна логика реализации.",
    ],
    "database": [
        "Проектирование базы данных выполнялось с учётом повторного использования шаблонов. Если каждый пользовательский сценарий хранить как полностью независимый набор текста, система быстро станет избыточной. Поэтому справочная часть хранится отдельно, а пользовательские сущности ссылаются на неё.",
        "Связь между задачами и документами позволяет строить полезные подсказки. Документ может быть нужен в нескольких задачах, поэтому при отображении списка документов важно убирать дубли и показывать назначение документа.",
        "Адреса пользователя выделены в отдельную логику, потому что один человек может иметь несколько значимых мест: дом, работа, адрес родственников. В дальнейшем это позволит показывать учреждения отдельно по каждому району.",
        "Новости и источники связаны с категориями. Такой подход позволяет показывать пользователю материалы по интересам, а также фильтровать ленту по официальным ресурсам. Для редактора это создаёт понятную модель публикации.",
        "Миграции базы данных фиксируют эволюцию проекта. Они показывают, что система развивалась поэтапно: от сценариев и пользователей к документам, статьям, просмотрам, адресам и заметкам. Это важный признак инженерной работы над проектом.",
    ],
    "architecture": [
        "Архитектурное решение учитывает текущий статус проекта. Ранний Flet-интерфейс был заменён web-интерфейсом, потому что сложный UX удобнее реализовывать в React. При этом Flet не исчез полностью: он используется как тонкая оболочка, решающая нативные задачи загрузки.",
        "Frontend и backend связаны через API-клиент. Такой слой нужен, чтобы страницы не зависели напрямую от реализации запросов. Если endpoint изменится, правка будет локализована в сервисном модуле, а не разбросана по компонентам.",
        "Store выполняет роль промежуточного слоя пользовательского состояния. Он объединяет данные из API, локального fallback и текущей сессии. Благодаря этому интерфейс остаётся устойчивым даже при частичной недоступности сервера.",
        "Backend-разделён на модели, схемы, API, сервисы, миграции и seed-данные. Такое разделение помогает поддерживать код и делает структуру проекта понятной для дальнейшей разработки.",
        "Административные функции не должны быть доступны обычному пользователю. Поэтому в интерфейсе и backend-логике предусматривается проверка роли. В будущем эта логика должна быть усилена полноценным RBAC и JWT-авторизацией.",
    ],
    "testing": [
        "Тестирование web-системы включает не только проверку отдельных кнопок, но и сквозные пользовательские сценарии. Для «Белпомощника» наиболее важен путь от выбора сценария до изменения прогресса. Если этот путь работает, основная идея проекта демонстрируется убедительно.",
        "Проверка ролей выполняется через быстрый вход под тестовыми пользователями. Это удобно для демонстрации руководителю: можно показать, что гость, гражданин, редактор и администратор видят разные возможности.",
        "Отдельно проверяется стабильность маршрутов. Пользователь не должен попадать на пустую страницу после обновления или прямого перехода по URL. Поэтому важные маршруты `/settings`, `/news`, `/law-detail`, `/profile`, `/documents` и `/situations` проверяются отдельно.",
        "Для WebView-оболочки важны не только успешные состояния. Если нет интернета или сервер недоступен, пользователь должен увидеть понятную страницу с кнопкой повторной попытки, а не технический текст.",
    ],
    "guide": [
        "Руководство пользователя должно быть написано не как инструкция для разработчика, а как описание реального пути гражданина. Поэтому в нём используются названия разделов интерфейса и действия, которые пользователь видит на экране.",
        "Основной сценарий демонстрации начинается с главной страницы, потому что именно она объясняет назначение продукта. Далее пользователь переходит в каталог и выбирает материал. Такой порядок показывает, что система не требует знания внутренней структуры.",
        "Работа с документами объясняется отдельно, поскольку это одна из наиболее чувствительных функций. Пользователь должен понимать, что документ может быть как требованием к задаче, так и личной записью со сроком действия.",
        "Административная панель в руководстве описывается кратко, но обязательно. Для дипломного проекта важно показать, что система не является статичным набором страниц. Она предполагает дальнейшее наполнение, модерацию и управление.",
        "При показе мобильной версии следует акцентировать внимание на том, что интерфейс остаётся тем же, но открывается внутри WebView-оболочки. Пользователь не видит браузерный интерфейс и получает более цельный опыт приложения.",
    ],
    "economics": [
        "Экономический эффект проекта имеет не только денежное выражение. Для гражданина важны снижение неопределённости, уменьшение количества повторных обращений и экономия личного времени. Эти факторы сложнее точно оценить, но они отражают реальную полезность системы.",
        "Для организации, которая могла бы использовать такую платформу, выгода выражается в снижении нагрузки на первичные консультации. Если типовые вопросы объяснены заранее и представлены в виде сценариев, специалист тратит меньше времени на повторение базовой информации.",
        "Затраты на разработку включают не только программирование. Значительную долю занимают анализ предметной области, проектирование структуры данных, подготовка сценариев, оформление источников, тестирование и документация. Это соответствует характеру информационной системы.",
        "Использование открытых технологий снижает прямые лицензионные затраты. React, Vite, FastAPI, SQLAlchemy и SQLite доступны без покупки коммерческих лицензий, что важно для учебного и MVP-этапа.",
        "Внедрение системы может быть поэтапным. Сначала платформа используется как справочный и демонстрационный продукт, затем подключаются production-база, редакторский процесс, уведомления и расширенный набор сценариев. Такой подход уменьшает первоначальные риски.",
    ],
    "labor": [
        "Организация рабочего места программиста влияет не только на безопасность, но и на качество разработки. При неудобной посадке и плохом освещении быстрее накапливается усталость, увеличивается вероятность ошибок и снижается внимательность при работе с кодом.",
        "Правильное размещение ПК должно учитывать не только монитор, но и периферийные устройства. Клавиатура, мышь, зарядные устройства и дополнительные мониторы не должны создавать лишних проводов в проходах или вынуждать пользователя принимать неудобное положение.",
        "Важным профилактическим мероприятием является режим перерывов. Кратковременное переключение внимания, упражнения для глаз и изменение положения тела помогают снизить статическую нагрузку. Это особенно актуально при подготовке дипломного проекта, когда работа за компьютером длится много часов подряд.",
        "Расследование несчастных случаев необходимо рассматривать как систему предотвращения повторных происшествий. Даже если случай кажется незначительным, его фиксация позволяет выявить недостатки организации рабочего места или инструктажа.",
        "Для рабочего места с ПК типовыми опасными факторами являются электрический ток, неудачное расположение кабелей, недостаточное освещение, перегрев оборудования и длительное зрительное напряжение. Каждый из этих факторов должен быть учтён при организации безопасной работы.",
    ],
}


INTRO = [
    f"Современный пользователь всё чаще ожидает, что бытовые и административные вопросы можно решить через понятный цифровой сервис. При этом значительная часть жизненных процедур остаётся сложной: необходимо определить последовательность действий, найти официальные источники, подготовить документы, обратиться в подходящее учреждение и не пропустить срок. Методические требования к дипломному проекту предполагают, что разработка должна быть не только программно реализована, но и обоснована с точки зрения предметной области, архитектуры, экономической эффективности и охраны труда {ref(1)}.",
    f"Оформление пояснительной записки выполняется с учётом требований к текстовым документам и правилам оформления отчётной документации {ref(2)}, {ref(3)}. Поэтому в работе уделяется внимание не только функциональности веб-системы, но и качеству проектных решений: структуре данных, ролям пользователей, пользовательскому пути, источникам информации, тестированию и сопровождению.",
    "Дипломный проект посвящён разработке информационной веб-системы «Белпомощник». Платформа ориентирована на граждан Республики Беларусь и предназначена для помощи при решении административно-бытовых задач: оформление документов, поиск порядка действий при жизненной ситуации, контроль сроков, работа с уведомлениями, просмотр закон-апдейтов и получение справочных сведений о государственных учреждениях.",
    "Ключевая идея проекта состоит в том, что пользователь получает не разрозненную статью, а маршрут действий. Быстрая проблема открывается как справочная карточка, жизненный сценарий превращается в личную ситуацию, а задачи, документы и прогресс позволяют двигаться по плану без потери контекста. Такой подход делает сервис полезным как для первичного информирования гражданина, так и для демонстрации будущего развития системы до полноценной городской или республиканской платформы.",
    "Цель дипломного проекта – разработка информационной веб-системы «Белпомощник», обеспечивающей пошаговую поддержку граждан при решении административно-бытовых задач. Для достижения цели необходимо изучить предметную область, сформировать требования, спроектировать структуру данных, реализовать пользовательский интерфейс и backend-основу, подготовить разделы для разных ролей, выполнить тестирование и оценить экономическую целесообразность.",
    "Практический результат проекта представляет собой веб-платформу на React/Vite с серверной частью FastAPI и SQLite. Дополнительно подготовлена WebView-оболочка, позволяющая показывать тот же интерфейс как мобильное приложение. Такое решение сохраняет единый пользовательский интерфейс и одновременно даёт основу для дальнейшей упаковки под разные платформы.",
]


def add_paragraphs(doc, paragraphs):
    for text in paragraphs:
        p(doc, text)


def introduction(doc):
    page_break(doc)
    section_heading(doc, "ВВЕДЕНИЕ")
    add_paragraphs(doc, INTRO)
    tasks = [
        "изучить предметную область административно-бытовых задач граждан;",
        "сформулировать требования к веб-приложению и пользовательским сценариям;",
        "обосновать выбор архитектуры, frontend-стека, backend-стека и СУБД;",
        "спроектировать функциональную модель и структуру базы данных;",
        "реализовать каталог проблем, жизненные сценарии, личные ситуации, документы, уведомления, новости, профиль и административную панель;",
        "проверить работоспособность основных функций и подготовить руководство пользователя;",
        "выполнить экономическое обоснование и рассмотреть вопросы охраны труда.",
    ]
    p(doc, "Для достижения поставленной цели в работе решаются следующие задачи:")
    for item in tasks:
        p(doc, f"– {item}")
    p(doc, "Объектом разработки является информационная веб-система для граждан. Предметом разработки является программная реализация платформы «Белпомощник», включающей пользовательские, редакторские и административные функции.")


ANALYTIC_PARAS = {
    "1.1": [
        "Предметная область проекта связана с повседневными ситуациями, в которых гражданину требуется получить понятную информацию о действиях, документах, сроках и организациях. К таким ситуациям относятся оформление паспорта, регистрация ребёнка, открытие индивидуального предпринимательства, жилищно-коммунальные вопросы, налоговые обязательства, семейные процедуры и взаимодействие с официальными источниками.",
        "Основная сложность предметной области состоит в распределённости информации. Пользователь может найти сведения на официальных порталах, сайтах ведомств, страницах местных исполнительных органов или в нормативных актах, однако ему всё равно необходимо самостоятельно собрать картину целиком. В результате простое бытовое действие превращается в цепочку поиска, проверки и уточнения.",
        "Для пользователя особенно важны простота языка и последовательность действий. Если сервис ограничивается только справочным текстом, он не решает проблему организации процесса. Поэтому в «Белпомощнике» применяется модель, где проблема, сценарий, задача, документ и учреждение связаны между собой. Пользователь видит не только описание, но и следующий конкретный шаг.",
        "В предметной области выделяются три уровня информации. Первый уровень – справочная проблема, например «потерял паспорт». Второй уровень – жизненный сценарий, например «рождение ребёнка» или «открытие ИП», включающий несколько этапов. Третий уровень – личная ситуация пользователя, где шаблон превращается в индивидуальный план с задачами и прогрессом.",
        "Важным требованием является аккуратная работа с источниками. Система не должна выдавать справочную информацию как окончательную юридическую консультацию. Поэтому в интерфейсе предусматривается блок официальных источников, а в содержании сценариев используется предупреждение о необходимости проверки актуальных требований на официальных ресурсах Республики Беларусь.",
        "Отдельное значение имеет доступность интерфейса. Целевая аудитория включает пользователей разного возраста, опыта и уровня цифровой грамотности. Поэтому структура страниц должна быть простой, кнопки – понятными, а важные действия – доступными без изучения сложных инструкций.",
    ],
    "1.2": [
        "Постановка задачи формулируется исходя из необходимости создать веб-систему, которая объединяет справочный каталог, личные планы и административное наполнение контента. Система должна помогать пользователю не только узнать информацию, но и применить её к своей ситуации.",
        "Пользователь должен иметь возможность открыть главную страницу, перейти в каталог, выбрать проблему или жизненный сценарий, изучить описание, этапы, документы и источники. После этого он может создать личную ситуацию, где задачи автоматически добавляются в список, а выполнение задач влияет на прогресс.",
        "Для личной работы в системе предусматриваются документы пользователя, уведомления, профиль, адреса, настройки доступности и предпочтения источников. Гость может просматривать основную информацию, но действия, связанные с персональными данными, должны требовать входа.",
        "Для поддержки жизненного цикла контента предусматриваются роли редактора и администратора. Редактор отвечает за проблемы, сценарии, статьи и официальные источники. Администратор управляет пользователями, ролями, справочниками и аналитикой.",
        "Разрабатываемая система должна иметь расширяемую архитектуру. На этапе дипломного проекта используется SQLite как MVP-хранилище, однако структура backend-слоя должна позволять в дальнейшем перейти к промышленной СУБД, полноценной авторизации, серверному хранению документов и push/email-уведомлениям.",
    ],
    "1.3": [
        "При анализе систем-аналогов рассматривались электронные государственные порталы, справочные сайты, базы нормативных актов и универсальные сервисы задач. Каждый из этих типов решений закрывает часть потребности, но не объединяет все элементы в единую пользовательскую траекторию.",
        "Официальные порталы предоставляют достоверную информацию и электронные услуги, однако их структура часто ориентирована на ведомственную логику. Пользователь должен знать, какой орган отвечает за конкретный вопрос, и самостоятельно понять, какие действия связаны между собой.",
        "Справочные сайты используют более простой язык, но не всегда дают пользователю персональный план. Они могут объяснить порядок действий, однако не сохраняют прогресс, не связывают документы с задачами и не дают роли администратора для системного наполнения контента.",
        "Универсальные планировщики задач позволяют отмечать выполнение, но не содержат предметных шаблонов, официальных источников и справочников учреждений. Для административно-бытовых задач такой инструмент становится слишком общим и требует ручного ввода всего сценария.",
        "Таким образом, «Белпомощник» занимает промежуточную позицию: он не заменяет официальные ресурсы, а связывает их с понятным пользовательским маршрутом. Это делает систему более прикладной для гражданина и более управляемой для администратора контента.",
    ],
    "1.4": [
        "Требования к веб-приложению разделяются на функциональные и нефункциональные. Функциональные требования описывают действия пользователя и администратора, а нефункциональные задают качество интерфейса, устойчивость, безопасность, адаптивность и сопровождаемость.",
        "К пользовательским функциям относятся просмотр главной страницы, поиск, каталог проблем и сценариев, создание личной ситуации, отметка задач, работа с документами, получение уведомлений, просмотр новостей и источников, редактирование профиля и использование настроек доступности.",
        "К редакторским функциям относится подготовка справочного контента: проблем, жизненных сценариев, этапов, задач, документов, источников и закон-апдейтов. Для администратора важны управление ролями, справочниками, пользователями, аналитикой и контроль качества данных.",
        "Нефункциональные требования включают адаптивность под desktop, tablet и mobile, понятные состояния ошибок, работу гостевого режима, защиту персональных действий, единое хранение пользовательского состояния и возможность дальнейшего подключения production-инфраструктуры.",
        "Так как система предназначена для широкого круга пользователей, интерфейс должен использовать спокойную визуальную подачу, крупную типографику, понятные подписи и последовательную навигацию. Сложные действия должны сопровождаться подсказками, а пустые состояния – объяснять следующий шаг.",
    ],
    "1.5": [
        f"Frontend-часть веб-системы построена на React. Эта библиотека широко используется для создания интерактивных пользовательских интерфейсов и хорошо подходит для компонентного подхода {ref(4)}. Сборка проекта выполняется с помощью Vite, что ускоряет локальную разработку и упрощает конфигурацию frontend-приложения {ref(5)}.",
        f"Для повышения надёжности frontend-кода используется TypeScript {ref(6)}. Типизация помогает фиксировать структуру данных проблем, сценариев, документов, профиля и пользовательских ситуаций. Это особенно важно для проекта, где одни и те же сущности передаются между интерфейсом, backend API и локальным fallback-хранилищем.",
        f"Backend-часть реализована на FastAPI {ref(7)}. Этот фреймворк выбран благодаря простому описанию маршрутов, удобной работе со схемами данных и хорошей пригодности для REST API. Для слоя доступа к данным используется SQLAlchemy {ref(8)}, что позволяет отделить модели и бизнес-логику от конкретной СУБД.",
        f"На этапе MVP в качестве СУБД применяется SQLite {ref(9)}. Такой выбор оправдан для дипломного проекта: база не требует отдельного сервера, быстро разворачивается локально и подходит для демонстрации структуры данных. При дальнейшем развитии проекта возможен переход к PostgreSQL без изменения пользовательской логики.",
        f"В качестве дополнительного направления учитывается возможность упаковки интерфейса в PWA и WebView. Подход PWA описан в открытой документации web-платформы {ref(10)}. Он позволяет использовать один интерфейс для браузера и мобильной оболочки, а WebView-слой отвечает за нативные состояния загрузки, ошибки сети и будущие функции приватности приложения.",
    ],
    "1.5.1": [
        "Архитектура системы построена по клиент-серверному принципу. Пользователь взаимодействует с React/Vite-интерфейсом, который обращается к backend API. Сервер отвечает за авторизацию, пользовательские данные, сценарии, документы, уведомления, статьи, учреждения и административные функции.",
        "Такое разделение позволяет не смешивать визуальный слой и бизнес-логику. Frontend отвечает за удобство, маршруты, формы, фильтры и отображение данных. Backend обеспечивает проверку прав, сохранение данных, миграции, seed-данные и общую модель предметной области.",
        "Дополнительным элементом является WebView-оболочка. Она не содержит собственного продуктового интерфейса, а только открывает React-сайт и показывает нативные состояния при загрузке или отсутствии сети. Это снижает стоимость сопровождения: изменения основного интерфейса автоматически доступны и в web-версии, и в мобильной оболочке.",
        f"В проекте учитываются базовые принципы безопасности web-приложений, включая разграничение ролей, защиту персональных действий и осторожную работу с пользовательскими данными. При дальнейшей доработке целесообразно использовать рекомендации OWASP по тестированию и защите web-систем {ref(11)}.",
    ],
    "1.5.2": [
        "SQLite выбрана как локальная реляционная СУБД для MVP. Она позволяет хранить связанные данные: пользователей, роли, профиль, адреса, сценарии, этапы, задачи, документы, учреждения, новости, источники и уведомления. Для дипломного проекта важна воспроизводимость: база может быть создана миграциями и заполнена seed-данными без сложной инфраструктуры.",
        "В дальнейшем система может быть переведена на PostgreSQL. Это потребуется при росте числа пользователей, появлении серверного хранения файлов, полноценной аналитики, регулярных уведомлений и редакторской работы нескольких сотрудников. Однако выбранная структура моделей уже готовит проект к такому переходу.",
        "Использование ORM снижает зависимость прикладного кода от конкретной базы. Модели описывают предметную область, а миграции фиксируют изменение структуры. Благодаря этому практический раздел проекта может показать не только интерфейс, но и основу будущей информационной системы.",
    ],
}


def analytical_section(doc):
    page_break(doc)
    section_heading(doc, "1 АНАЛИТИЧЕСКИЙ РАЗДЕЛ")
    subheading(doc, "1.1 Описание предметной области")
    add_paragraphs(doc, ANALYTIC_PARAS["1.1"])
    p(doc, "В таблице 1.1 представлено распределение основных групп административно-бытовых задач, которые учитываются при проектировании системы.")
    add_table_caption(doc, "Таблица 1.1 – Группы административно-бытовых задач")
    add_doc_table(doc, ["Группа задач", "Примеры", "Особенности обработки в системе"], [
        ["Документы", "паспорт, ID-карта, медкнижка, свидетельства", "контроль сроков, реквизиты, связь с задачами"],
        ["Семья", "рождение ребёнка, брак, развод", "жизненные сценарии с этапами и документами"],
        ["Жильё и ЖКХ", "регистрация, лицевые счета, оплаты", "связь с адресами и учреждениями"],
        ["Налоги и ИП", "открытие ИП, обязательства, сроки", "персональные напоминания и справочные материалы"],
        ["Правовая информация", "закон-апдейты, источники", "лента новостей и фильтрация по интересам"],
    ], [3.2, 5.3, 7.0])
    add_paragraphs(doc, [
        "Такое распределение показывает, что платформа должна работать не как обычный каталог статей, а как система связанных сущностей. Каждая группа задач требует своего набора данных, но пользователь должен видеть их в едином интерфейсе.",
        "Для жизненных сценариев особенно важно фиксировать этапность. Например, при рождении ребёнка сначала появляются медицинские документы, затем регистрация рождения, получение свидетельства, оформление пособий, регистрация по месту жительства и прикрепление к поликлинике. Если эти шаги не связаны, пользователь теряет последовательность.",
    ])
    expand(doc, "domain", EXTRA["domain"])
    subheading(doc, "1.2 Постановка задачи")
    add_paragraphs(doc, ANALYTIC_PARAS["1.2"])
    p(doc, "В таблице 1.2 представлены основные задачи разработки, которые вытекают из постановки дипломного проекта.")
    add_table_caption(doc, "Таблица 1.2 – Основные задачи разработки")
    add_doc_table(doc, ["№", "Задача", "Ожидаемый результат"], [
        ["1", "Создать каталог проблем и сценариев", "Пользователь быстро находит подходящий материал"],
        ["2", "Реализовать личные ситуации", "Шаблон превращается в персональный план"],
        ["3", "Связать задачи и документы", "Пользователь понимает, что подготовить к каждому шагу"],
        ["4", "Добавить роли", "Гость, гражданин, редактор и администратор имеют разные права"],
        ["5", "Подготовить backend API", "Интерфейс может работать с серверными данными"],
        ["6", "Обеспечить адаптивность", "Сервис пригоден для desktop, tablet и mobile"],
    ], [1.0, 6.0, 8.5])
    expand(doc, "task", EXTRA["task"])
    subheading(doc, "1.3 Обзор систем-аналогов")
    add_paragraphs(doc, ANALYTIC_PARAS["1.3"])
    p(doc, "В таблице 1.3 приведено сравнение «Белпомощника» с типовыми группами аналогов.")
    add_table_caption(doc, "Таблица 1.3 – Сравнение с системами-аналогами")
    add_doc_table(doc, ["Критерий", "Официальный портал", "Справочный сайт", "Планировщик задач", "Белпомощник"], [
        ["Официальность источников", "высокая", "средняя", "не применимо", "через ссылки на источники"],
        ["Пошаговый план", "частично", "частично", "ручной ввод", "есть"],
        ["Личная ситуация", "обычно нет", "нет", "есть, но без предметной логики", "есть"],
        ["Документы внутри задач", "частично", "частично", "нет", "есть"],
        ["Роли редактора и администратора", "внутренние", "зависит от сайта", "нет", "предусмотрены"],
        ["Адаптация под мобильный сценарий", "разная", "разная", "обычно есть", "есть web и WebView"],
    ], [3.0, 3.3, 3.5, 3.8, 3.4])
    add_paragraphs(doc, [
        "Сравнение показывает, что ни один тип аналога полностью не закрывает задачу проекта. Поэтому разработка собственной системы оправдана: она объединяет справочную информацию, личные планы, документы, уведомления и роли управления контентом.",
        "При этом «Белпомощник» не конкурирует с официальными порталами как источник нормативной информации. Его назначение – сделать путь пользователя более понятным и направить его к правильным источникам, учреждениям и действиям.",
    ])
    expand(doc, "analogs", EXTRA["analogs"])
    subheading(doc, "1.4 Анализ требований к веб-приложению")
    add_paragraphs(doc, ANALYTIC_PARAS["1.4"])
    p(doc, "В таблице 1.4 представлены функциональные требования к основным ролям системы.")
    add_table_caption(doc, "Таблица 1.4 – Функциональные требования по ролям")
    add_doc_table(doc, ["Роль", "Доступные действия", "Ограничения"], [
        ["Гость", "просмотр главной, каталога, новостей и справочных материалов", "не может создавать личные данные"],
        ["Гражданин", "создание ситуаций, отметка задач, документы, профиль, уведомления", "не имеет доступа к админ-функциям"],
        ["Редактор", "работа с проблемами, сценариями, статьями и источниками", "не управляет системными ролями"],
        ["Администратор", "пользователи, роли, справочники, аналитика, публикации", "доступ ограничен служебной ролью"],
    ], [3.0, 7.4, 5.0])
    add_paragraphs(doc, [
        "Помимо ролей учитываются требования к адаптивности. На desktop пользователь чаще работает с широкими таблицами, формами и админ-панелью. На мобильном устройстве важнее компактный путь: нижняя навигация, крупные зоны нажатия, отсутствие горизонтального скролла и понятные состояния загрузки.",
        "Нефункциональные требования также включают устойчивость к отсутствию backend. В проекте сохраняется mock fallback, который позволяет интерфейсу не падать при недоступном API. Для демонстрационного этапа это особенно важно, поскольку сервис должен открываться в web и в WebView-оболочке.",
    ])
    expand(doc, "requirements", EXTRA["requirements"])
    subheading(doc, "1.5 Обоснование выбора технологического стека")
    add_paragraphs(doc, ANALYTIC_PARAS["1.5"])
    p(doc, "В таблице 1.5 представлено назначение основных технологий, использованных в проекте.")
    add_table_caption(doc, "Таблица 1.5 – Технологический стек проекта")
    add_doc_table(doc, ["Компонент", "Технология", "Назначение"], [
        ["Frontend", "React, Vite, TypeScript", "пользовательский интерфейс, маршруты, формы, адаптивность"],
        ["Backend", "FastAPI", "REST API, роли, пользовательские данные, админ-функции"],
        ["База данных", "SQLite", "MVP-хранение сущностей и миграции"],
        ["ORM", "SQLAlchemy", "модели, связи, переносимость между СУБД"],
        ["Мобильная оболочка", "Flet WebView", "загрузка React-сайта как приложения"],
    ], [3.0, 4.2, 8.2])
    expand(doc, "stack", EXTRA["stack"])
    subheading(doc, "1.5.1 Обоснование выбора архитектуры и сервера", 3)
    add_paragraphs(doc, ANALYTIC_PARAS["1.5.1"])
    subheading(doc, "1.5.2 Обоснование выбора СУБД", 3)
    add_paragraphs(doc, ANALYTIC_PARAS["1.5.2"])


def practical_section(doc):
    page_break(doc)
    section_heading(doc, "2 ПРАКТИЧЕСКИЙ РАЗДЕЛ")
    subheading(doc, "2.1 Функциональная модель")
    add_paragraphs(doc, [
        "Функциональная модель описывает взаимодействие пользователя с системой и показывает, какие процессы должны быть поддержаны программным продуктом. Для «Белпомощника» центральным процессом является переход от справочной информации к персональному плану действий.",
        "В модели выделяются четыре группы участников: гость, гражданин, редактор контента и администратор. Гость знакомится с возможностями сервиса, гражданин создаёт личные данные, редактор наполняет контент, администратор управляет системными настройками и справочниками.",
        "Функциональная модель важна для проектирования интерфейса. Если пользователь не понимает, чем проблема отличается от жизненного сценария или личной ситуации, система теряет смысл. Поэтому маршруты, подписи, фильтры и карточки должны закреплять эту терминологию.",
    ])
    expand(doc, "functional", EXTRA["functional"])
    subheading(doc, "2.1.1 Диаграмма вариантов использования", 3)
    p(doc, "На рисунке 2.1 показана диаграмма вариантов использования платформы «Белпомощник». Она отражает основные роли и действия, которые доступны каждой роли.")
    add_picture(doc, IMG_DIR / "use_case.png", 15.2)
    add_figure_caption(doc, "Рисунок 2.1 – Диаграмма вариантов использования")
    add_paragraphs(doc, [
        "Гражданин является основной ролью системы. Для него предусмотрены поиск, просмотр проблем, выбор жизненных сценариев, создание личной ситуации, отметка задач, добавление документов и получение уведомлений.",
        "Редактор контента отвечает за качество материалов. Он создаёт и редактирует проблемы, сценарии, этапы, задачи, статьи, закон-апдейты и источники. Администратор дополняет эти возможности управлением пользователями, ролями, справочниками и аналитикой.",
    ])
    subheading(doc, "2.1.2 Диаграмма активности", 3)
    p(doc, "На рисунке 2.2 представлена диаграмма активности, описывающая основной пользовательский путь от открытия главной страницы до автоматического пересчёта прогресса.")
    add_picture(doc, IMG_DIR / "activity.png", 15.2)
    add_figure_caption(doc, "Рисунок 2.2 – Диаграмма активности создания личной ситуации")
    add_paragraphs(doc, [
        "Процесс начинается с выбора проблемы или категории. Пользователь изучает материал, после чего принимает решение создать личную ситуацию. Система формирует задачи и документы на основе шаблона, а дальнейшее выполнение задач отражается в прогрессе.",
        "Такой процесс уменьшает количество ручных действий. Пользователю не нужно самостоятельно переносить пункты из статьи в список задач: шаблон уже содержит этапы, документы, сроки и связи.",
    ])
    subheading(doc, "2.1.3 Диаграмма последовательности", 3)
    p(doc, "На рисунке 2.3 показана диаграмма последовательности для операции создания личной ситуации из шаблона жизненного сценария.")
    add_picture(doc, IMG_DIR / "sequence.png", 15.2)
    add_figure_caption(doc, "Рисунок 2.3 – Диаграмма последовательности создания ситуации")
    add_paragraphs(doc, [
        "Диаграмма показывает, что интерфейс не создаёт данные изолированно. React-компонент вызывает действие store/API-клиента, затем при наличии авторизованного backend-соединения отправляется запрос на сервер. Сервер сохраняет ситуацию и связанные задачи в базе данных, после чего интерфейс обновляет локальное состояние.",
        "При недоступности backend сохраняется локальный fallback. Это решение важно для демонстрационного этапа, но архитектура уже предусматривает переход к полному серверному хранению.",
    ])
    subheading(doc, "2.2 Проектирование структуры базы данных")
    add_paragraphs(doc, [
        "Структура базы данных строится вокруг сценариев и пользовательских данных. Шаблон сценария содержит этапы и задачи, а пользовательская ситуация связывает конкретного пользователя с выбранным шаблоном и состоянием выполнения.",
        "В базе также учитываются документы, профиль, адреса, уведомления, статьи, источники, учреждения и роли. Такая структура позволяет развивать систему постепенно: сначала хранить MVP-данные, затем добавлять расширенную персонализацию, push-уведомления и редакторский workflow.",
    ])
    p(doc, "На рисунке 2.4 представлена схема основных сущностей базы данных и связей между ними.")
    add_picture(doc, IMG_DIR / "database.png", 15.2)
    add_figure_caption(doc, "Рисунок 2.4 – Схема основных сущностей базы данных")
    p(doc, "В таблице 2.1 приведено назначение ключевых сущностей базы данных.")
    add_table_caption(doc, "Таблица 2.1 – Основные сущности базы данных")
    add_doc_table(doc, ["Сущность", "Назначение", "Основные связи"], [
        ["users", "хранение аккаунтов и ролей", "profile, situations, documents"],
        ["scenario_templates", "шаблоны жизненных сценариев", "stages, sources, authorities"],
        ["scenario_tasks", "задачи внутри этапов", "documents, dependencies"],
        ["user_situations", "личные ситуации пользователя", "user, template, tasks"],
        ["user_tasks", "состояние выполнения задач", "situation, source task"],
        ["documents", "личные и справочные документы", "user, task, expiry"],
        ["articles", "новости и закон-апдейты", "sources, categories"],
        ["authorities", "государственные учреждения", "region, city, service type"],
    ], [3.7, 6.2, 5.7])
    add_paragraphs(doc, [
        "В структуре данных важно не дублировать справочную и пользовательскую информацию. Шаблон хранит описание типового сценария, а пользовательская ситуация хранит только то, что относится к конкретному человеку: прогресс, выполненные задачи, заметки и персональные сроки.",
        "Для документов предусмотрено разделение справочных требований и личных документов. Справочный документ указывает, что потребуется для выполнения задачи, а личный документ содержит реквизиты и срок действия, относящиеся к пользователю.",
    ])
    expand(doc, "database", EXTRA["database"])
    subheading(doc, "2.3 Архитектура программного решения")
    add_paragraphs(doc, [
        "Архитектура проекта отражает текущий переход от ранней Flet-версии к web-платформе. Основной интерфейс реализован в каталоге React/Vite. Backend расположен отдельно и предоставляет API для публичных данных, пользовательских ситуаций, документов, уведомлений, профиля и администрирования.",
        "В системе используется принцип API-first. Если backend доступен и пользователь авторизован, данные синхронизируются с сервером. Если API временно недоступен, интерфейс сохраняет работоспособность через локальные данные и mock fallback. Это повышает устойчивость демонстрационной версии.",
    ])
    p(doc, "На рисунке 2.5 показана архитектура программного решения и взаимодействие frontend, backend, базы данных и мобильной WebView-оболочки.")
    add_picture(doc, IMG_DIR / "architecture.png", 15.2)
    add_figure_caption(doc, "Рисунок 2.5 – Архитектура программного решения")
    p(doc, "В таблице 2.2 приведено распределение ответственности между основными подсистемами.")
    add_table_caption(doc, "Таблица 2.2 – Ответственность подсистем")
    add_doc_table(doc, ["Подсистема", "Ответственность", "Примеры функций"], [
        ["React frontend", "маршруты, формы, отображение данных", "каталог, профиль, документы, новости"],
        ["Store/API client", "единая работа с local fallback и API", "создание ситуации, синхронизация документов"],
        ["FastAPI backend", "проверка прав и обработка запросов", "auth, CRUD, админ-эндпоинты"],
        ["SQLite", "хранение MVP-данных", "пользователи, сценарии, задачи, статьи"],
        ["WebView shell", "нативная загрузка сайта", "offline, retry, portrait orientation"],
    ], [3.2, 6.3, 6.1])
    add_paragraphs(doc, [
        "С точки зрения сопровождения важно, что WebView-оболочка не дублирует интерфейс. Она отвечает только за нативные состояния, а весь пользовательский опыт сосредоточен в React-приложении. Это уменьшает риск расхождения web и мобильной версии.",
        "Административная часть реализуется как часть web-интерфейса, но закрывается role guard. Обычный гражданин не видит полный админ-раздел, редактор получает доступ к контентным функциям, а администратор работает с пользователями, ролями и справочниками.",
    ])
    expand(doc, "architecture", EXTRA["architecture"])
    subheading(doc, "2.4 Тестирование программного продукта")
    add_paragraphs(doc, [
        "Тестирование программного продукта выполнялось по основным пользовательским и административным сценариям. Для дипломного проекта важно подтвердить, что пользователь может пройти путь от выбора сценария до контроля прогресса, а также что роли и защищённые действия работают предсказуемо.",
        "Дополнительно проверялась сборка frontend-приложения, запуск backend, стабильность маршрутов, открытие основных страниц, работа форм, состояние гостя и переключение тестовых пользователей.",
    ])
    p(doc, "В таблице 2.3 представлены основные тест-кейсы, использованные для проверки веб-системы.")
    add_table_caption(doc, "Таблица 2.3 – Результаты тестирования программного продукта")
    add_doc_table(doc, ["№", "Проверка", "Ожидаемый результат", "Результат"], [
        ["1", "Открытие главной страницы", "дашборд отображается без ошибок", "выполнено"],
        ["2", "Переход в каталог", "отображаются проблемы и сценарии", "выполнено"],
        ["3", "Создание ситуации", "задачи добавляются из шаблона", "выполнено"],
        ["4", "Отметка задачи", "прогресс пересчитывается", "выполнено"],
        ["5", "Добавление документа", "документ появляется в списке", "выполнено"],
        ["6", "Гостевой режим", "личные действия требуют вход", "выполнено"],
        ["7", "Админ-панель", "доступ только для администратора", "выполнено"],
        ["8", "Сборка frontend", "pnpm build завершается без ошибок", "выполнено"],
    ], [0.8, 4.9, 6.2, 3.0])
    add_paragraphs(doc, [
        "Результаты проверки показывают, что первая web-версия проекта пригодна для демонстрации основного пользовательского пути. При этом часть возможностей остаётся направлением дальнейшей разработки: production-авторизация, серверное хранение файлов, расширенная аналитика, реальные push-уведомления и юридически проверенное наполнение контента.",
        "Отдельное тестирование требуется для мобильной WebView-оболочки на реальных устройствах. В рамках проекта уже подготовлены нативные состояния загрузки, ошибки сети и повторной попытки, однако финальная проверка должна выполняться на фактическом Android/iOS-устройстве.",
    ])
    expand(doc, "testing", EXTRA["testing"])


def user_guide(doc):
    page_break(doc)
    section_heading(doc, "3 РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ")
    add_paragraphs(doc, [
        "Работа с веб-платформой начинается с главной страницы. Пользователь видит основные направления помощи, быстрый переход в каталог, личные ситуации, документы, новости и профиль. Интерфейс построен так, чтобы пользователь мог начать с проблемы, категории или готового жизненного сценария.",
        "Гость может просматривать публичные страницы, но при попытке создать личную ситуацию, добавить документ или изменить профиль система предлагает войти или зарегистрироваться. Это разделение позволяет показать возможности продукта без раскрытия персональных действий.",
    ])
    p(doc, "На рисунке 3.1 показана главная страница веб-платформы «Белпомощник».")
    add_picture(doc, IMG_DIR / "screen_home.png", 15.2)
    add_figure_caption(doc, "Рисунок 3.1 – Главная страница веб-платформы")
    add_paragraphs(doc, [
        "Для поиска нужной темы пользователь открывает каталог. В каталоге отдельно выделяются проблемы и жизненные сценарии. Проблема используется для быстрой справки, а жизненный сценарий – для создания персонального плана.",
    ])
    p(doc, "На рисунке 3.2 представлен каталог проблем и жизненных сценариев.")
    add_picture(doc, IMG_DIR / "screen_catalog.png", 15.2)
    add_figure_caption(doc, "Рисунок 3.2 – Каталог проблем и жизненных сценариев")
    add_paragraphs(doc, [
        "После выбора жизненного сценария пользователь просматривает описание, этапы, задачи, документы, учреждения и официальные источники. Кнопка создания личной ситуации переносит шаблон в раздел «Мои ситуации».",
        "В разделе «Мои ситуации» отображаются созданные пользователем планы. В карточке ситуации показываются статус, прогресс, количество задач и ближайшие действия. Выполненные задачи отмечаются пользователем вручную, а прогресс пересчитывается автоматически.",
    ])
    p(doc, "На рисунке 3.3 показан раздел личных ситуаций пользователя.")
    add_picture(doc, IMG_DIR / "screen_situations.png", 15.2)
    add_figure_caption(doc, "Рисунок 3.3 – Раздел «Мои ситуации»")
    add_paragraphs(doc, [
        "Документы в системе используются в двух смыслах. Справочные документы показывают, что потребуется для выполнения задачи. Личные документы позволяют пользователю хранить реквизиты, сроки действия и данные о файле скана. По умолчанию чувствительные данные должны отображаться осторожно.",
    ])
    p(doc, "На рисунке 3.4 представлен раздел документов.")
    add_picture(doc, IMG_DIR / "screen_documents.png", 15.2)
    add_figure_caption(doc, "Рисунок 3.4 – Раздел «Документы»")
    add_paragraphs(doc, [
        "Новости и закон-апдейты используются для проактивного информирования. Пользователь может просматривать материалы, связанные с документами, налогами, семьёй, работой и другими категориями. Источники информации должны быть связаны с публикациями и отображаться как справочная основа материала.",
        "Профиль пользователя содержит персональные сведения, адреса, настройки интерфейса, предпочтения источников и дополнительные элементы персонализации. Адреса используются для будущего подбора ближайших учреждений по району или населенному пункту.",
        "Администратор и редактор работают с отдельными служебными разделами. Они предназначены для управления сценариями, проблемами, статьями, источниками, справочниками, пользователями и ролями. В обычном пользовательском режиме эти функции скрыты.",
    ])
    p(doc, "На рисунке 3.5 показан общий вид административной панели.")
    add_picture(doc, IMG_DIR / "screen_admin.png", 15.2)
    add_figure_caption(doc, "Рисунок 3.5 – Административная панель")
    add_paragraphs(doc, [
        "Для демонстрации проекта рекомендуется использовать следующий сценарий: открыть главную страницу, перейти в каталог, выбрать жизненный сценарий «Рождение ребёнка», просмотреть этапы и документы, создать личную ситуацию, отметить несколько задач выполненными, затем вернуться на главную и показать изменение прогресса.",
        "Такой путь показывает основную ценность системы: пользователь получает не только справочную информацию, но и управляемый план действий. Дополнительно можно показать документы, новости, профиль, настройки и административные возможности.",
    ])
    expand(doc, "guide", EXTRA["guide"])


def economics(doc):
    page_break(doc)
    section_heading(doc, "4 ЭКОНОМИЧЕСКИЙ РАЗДЕЛ")
    paras = [
        "Экономический раздел предназначен для оценки затрат на создание программного продукта и определения целесообразности его внедрения. Для веб-системы «Белпомощник» экономический эффект связан с сокращением времени на поиск информации, уменьшением количества ошибок при подготовке документов и снижением нагрузки на первичные консультации.",
        "Расчёты выполняются для учебного программного продукта, поэтому часть исходных данных принимается условно. При этом структура расчёта соответствует методическим рекомендациям по экономическому разделу дипломных проектов: определяется трудоёмкость, заработная плата разработчика, отчисления, накладные расходы, прогнозируемая цена и возможный эффект внедрения.",
        "Разработка выполняется одним учащимся с использованием персонального компьютера, открытых инструментов разработки и локальной инфраструктуры. Основные затраты формируются за счёт трудоёмкости проектирования, программирования, тестирования и оформления материалов.",
    ]
    add_paragraphs(doc, paras)
    p(doc, "В таблице 4.1 представлена оценка трудоёмкости этапов разработки программного продукта.")
    add_table_caption(doc, "Таблица 4.1 – Расчёт трудоёмкости разработки")
    add_doc_table(doc, ["Этап", "Содержание работ", "Трудоёмкость, ч"], [
        ["Анализ", "изучение предметной области и требований", "48"],
        ["Проектирование", "структура данных, диаграммы, архитектура", "72"],
        ["Frontend", "страницы, маршруты, адаптивный интерфейс", "150"],
        ["Backend", "API, модели, миграции, seed-данные", "110"],
        ["Тестирование", "проверка сценариев и исправление ошибок", "60"],
        ["Документация", "пояснительная записка и материалы проекта", "70"],
    ], [3.2, 8.0, 3.0])
    add_paragraphs(doc, [
        "Общая трудоёмкость разработки составляет 510 часов. Для расчёта заработной платы принимается условная часовая ставка 12,55 руб. Основная заработная плата определяется по формуле:",
        "Зосн = T × Cч, где T – трудоёмкость разработки, Cч – часовая ставка разработчика. Следовательно, Зосн = 510 × 12,55 = 6400,50 руб. (4.1)",
        "Дополнительная заработная плата принимается в размере 10 % от основной заработной платы:",
        "Здоп = 6400,50 × 0,10 = 640,05 руб. (4.2)",
        "Отчисления на социальные нужды рассчитываются от суммы основной и дополнительной заработной платы. При коэффициенте 34,6 % получаем:",
        "Осн = (6400,50 + 640,05) × 0,346 = 2436,03 руб. (4.3)",
        "Накладные расходы принимаются в размере 20 % от основной заработной платы:",
        "Рнакл = 6400,50 × 0,20 = 1280,10 руб. (4.4)",
    ])
    p(doc, "В таблице 4.2 представлен сводный расчёт затрат на разработку.")
    add_table_caption(doc, "Таблица 4.2 – Сводный расчёт затрат")
    add_doc_table(doc, ["Показатель", "Расчёт", "Сумма, руб."], [
        ["Основная заработная плата", "510 × 12,55", "6400,50"],
        ["Дополнительная заработная плата", "10 %", "640,05"],
        ["Отчисления", "34,6 %", "2436,03"],
        ["Накладные расходы", "20 %", "1280,10"],
        ["Материальные расходы", "условно", "160,00"],
        ["Итого себестоимость", "сумма затрат", "10916,68"],
    ], [4.9, 4.8, 3.0])
    add_paragraphs(doc, [
        "Себестоимость разработки программного продукта составляет 10916,68 руб. Для расчёта прогнозируемой цены учитывается плановая прибыль. При норме прибыли 20 % цена программного средства составит:",
        "П = 10916,68 × 0,20 = 2183,34 руб. (4.5)",
        "Ц = 10916,68 + 2183,34 = 13100,02 руб. (4.6)",
        "Экономический эффект от внедрения связан с сокращением времени, которое пользователь или консультант тратит на первичный поиск информации. Если система позволяет сократить 380 часов в год при условной стоимости часа 12,55 руб., годовая экономия фонда времени составит:",
        "Эгод = 380 × 12,55 × 1,37 × 1,346 = 8791,63 руб. (4.7)",
        "Для оценки окупаемости принимается объём капитальных вложений, равный прогнозируемой цене программного продукта. Срок окупаемости рассчитывается по формуле:",
        "Ток = КВ / Эгод = 13100,02 / 8791,63 = 1,49 года. (4.8)",
        "Полученный срок окупаемости является приемлемым для информационной системы, которая может использоваться в образовательных, консультационных или муниципальных целях. При расширении аудитории и подключении реальных уведомлений эффект может увеличиться.",
    ])
    p(doc, "В таблице 4.3 приведены основные технико-экономические показатели мероприятия.")
    add_table_caption(doc, "Таблица 4.3 – Технико-экономические показатели")
    add_doc_table(doc, ["Показатель", "Значение"], [
        ["Общая трудоёмкость разработки", "510 ч"],
        ["Себестоимость разработки", "10916,68 руб."],
        ["Прогнозируемая цена", "13100,02 руб."],
        ["Годовая экономия", "8791,63 руб."],
        ["Срок окупаемости", "1,49 года"],
        ["Ожидаемый рост производительности информационной работы", "15–20 %"],
    ], [8.2, 5.8])
    add_paragraphs(doc, [
        "Разработка платформы «Белпомощник» экономически оправдана, поскольку продукт снижает затраты времени на поиск информации и уменьшает вероятность ошибок при подготовке документов. Наиболее заметный эффект может быть получен при использовании системы в качестве первичного навигатора по жизненным ситуациям.",
        "Экономический расчёт показывает, что даже при осторожных исходных данных проект имеет практическую ценность. Кроме прямой экономии времени, система повышает качество информационного обслуживания, обеспечивает единый подход к сценариям и создаёт основу для дальнейшей автоматизации уведомлений.",
    ])
    expand(doc, "economics", EXTRA["economics"])


def labor(doc):
    page_break(doc)
    section_heading(doc, "5 РАЗДЕЛ ПО ОХРАНЕ ТРУДА")
    subheading(doc, "5.1 Требования к помещениям и размещению ПК")
    add_paragraphs(doc, [
        f"Разработка веб-системы выполняется с использованием персонального компьютера, монитора, клавиатуры, мыши и сетевого оборудования. Работа программиста связана с длительной зрительной нагрузкой, статическим положением тела и необходимостью соблюдать требования безопасности при эксплуатации электрооборудования. Вопросы охраны труда рассматриваются на основе учебных и нормативных материалов {ref(20)}, {ref(21)}.",
        "Помещение, в котором размещается рабочее место разработчика, должно обеспечивать безопасные и комфортные условия труда. Рабочая зона должна иметь достаточную площадь, возможность проветривания, устойчивое размещение мебели и оборудования, свободный проход к рабочему месту и отсутствие предметов, создающих риск спотыкания.",
        "Персональный компьютер размещается на устойчивом рабочем столе. Монитор устанавливается так, чтобы верхняя граница экрана находилась примерно на уровне глаз или немного ниже. Расстояние от глаз до экрана должно быть удобным для чтения и не вызывать необходимости наклоняться вперёд.",
        "Клавиатура и мышь размещаются на поверхности стола таким образом, чтобы руки пользователя находились в естественном положении. Запястья не должны быть чрезмерно согнуты, плечи не должны подниматься, а локти должны располагаться близко к туловищу. Рабочее кресло должно обеспечивать поддержку спины и возможность регулировки высоты.",
        "Провода питания и соединительные кабели необходимо размещать так, чтобы они не пересекали проходы и не подвергались механическому повреждению. Не допускается эксплуатация оборудования с повреждённой изоляцией, нестабильным подключением, признаками перегрева или посторонним запахом.",
        "В помещении должна поддерживаться нормальная освещённость. Свет не должен создавать бликов на экране, а контраст между монитором и окружающей средой должен быть комфортным для зрения. При необходимости используются жалюзи, шторы или изменение положения монитора.",
    ])
    subheading(doc, "5.2 Расследование и учет несчастных случаев на производстве")
    add_paragraphs(doc, [
        "Несчастный случай на производстве представляет собой событие, в результате которого работник получил повреждение здоровья при выполнении трудовых обязанностей. Даже при работе за персональным компьютером полностью исключить риски нельзя: возможны падения, поражение электрическим током, травмы при перемещении оборудования или ухудшение состояния из-за нарушения режима труда.",
        "При возникновении несчастного случая необходимо немедленно прекратить работу, оказать первую помощь пострадавшему, при необходимости вызвать медицинскую помощь и сообщить ответственному лицу. Если ситуация связана с электрооборудованием, питание отключают только при условии, что это можно сделать безопасно.",
        "Расследование несчастного случая проводится для установления обстоятельств и причин происшествия. В процессе фиксируются место, время, характер работ, состояние оборудования, соблюдение инструкций, наличие опасных факторов и действия работников. Цель расследования состоит не только в оформлении документов, но и в предотвращении повторения аналогичных ситуаций.",
        "Учет несчастных случаев позволяет организации анализировать риски и принимать профилактические меры. Для рабочего места программиста такими мерами являются регулярный инструктаж, проверка исправности оборудования, организация перерывов, поддержание порядка на рабочем месте и соблюдение требований электробезопасности.",
        "В контексте дипломного проекта соблюдение требований охраны труда важно как при разработке программного продукта, так и при его дальнейшем использовании. Пользователь системы также взаимодействует с компьютером или мобильным устройством, поэтому интерфейс должен быть читаемым, не перегруженным и не требующим длительного напряжённого поиска нужного действия.",
    ])
    expand(doc, "labor", EXTRA["labor"])


def conclusion_sources_appendix(doc):
    page_break(doc)
    section_heading(doc, "ЗАКЛЮЧЕНИЕ")
    add_paragraphs(doc, [
        "В ходе выполнения дипломного проекта была разработана информационная веб-система «Белпомощник», предназначенная для поддержки граждан Республики Беларусь при решении административно-бытовых задач. Система объединяет каталог проблем, жизненные сценарии, личные ситуации, документы, уведомления, новости, профиль пользователя и административное управление контентом.",
        "В аналитическом разделе рассмотрена предметная область, показана проблема разрозненности справочной информации и обоснована необходимость сервиса, который превращает информацию в пошаговый план. Выполнен обзор аналогов, сформулированы требования к ролям, интерфейсу, данным, источникам и адаптивности.",
        "В практическом разделе спроектированы функциональная модель, диаграммы, структура базы данных и архитектура программного решения. Основной интерфейс реализован на React/Vite, серверная часть – на FastAPI, а MVP-хранилище – на SQLite. Дополнительно подготовлена WebView-оболочка, позволяющая демонстрировать систему как мобильное приложение.",
        "Руководство пользователя показывает основной путь работы: открыть главную страницу, выбрать проблему или жизненный сценарий, создать личную ситуацию, выполнять задачи и контролировать прогресс. Административная часть демонстрирует направление дальнейшего наполнения системы сценариями, статьями, источниками, пользователями и ролями.",
        "Экономический расчёт подтвердил целесообразность разработки программного продукта. Система сокращает время на поиск информации, снижает количество ошибок и повышает качество первичной ориентации гражданина в административно-бытовой ситуации. Раздел охраны труда показал необходимость правильной организации рабочего места и соблюдения порядка расследования несчастных случаев.",
        "Поставленная цель дипломного проекта достигнута. Разработана веб-платформа, демонстрирующая полный пользовательский путь и подготовленная к дальнейшему развитию. Следующими этапами могут стать подключение production-базы данных, расширение RBAC, серверное хранение файлов, реальные push/email-уведомления, юридически проверенное наполнение сценариев и публикация приложения.",
    ])
    page_break(doc)
    section_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    for i, source in enumerate(SOURCES, start=1):
        p(doc, f"{i}. {source}", "Diploma Source")
    page_break(doc)
    section_heading(doc, "ПРИЛОЖЕНИЕ А")
    p(doc, "Структура проекта «Белпомощник»", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_indent=False)
    add_paragraphs(doc, [
        "В приложении А приведено описание структуры программного проекта. Основной пользовательский интерфейс находится в каталоге reactvitemaket и представляет собой Vite/React-приложение. В нём размещены маршруты, страницы, компоненты, store, типы данных, сервисы поиска, API-клиент и стили.",
        "Backend-часть расположена в каталоге src/backend. Она включает FastAPI-приложение, модели SQLAlchemy, схемы, маршруты API, сервисный слой, миграции и seed-данные. Такое разделение позволяет поддерживать серверную часть независимо от пользовательского интерфейса.",
        "Файл src/mobile_webview.py является тонкой WebView-оболочкой. Он не содержит собственного интерфейса платформы, а открывает React-сайт и отвечает за нативные состояния загрузки, ошибки сети и повторной попытки.",
    ])
    p(doc, "В таблице А.1 представлена укрупнённая структура каталогов проекта.")
    add_table_caption(doc, "Таблица А.1 – Структура каталогов проекта")
    add_doc_table(doc, ["Каталог/файл", "Назначение"], [
        ["reactvitemaket/src/app", "основные страницы, маршруты и компоненты React-интерфейса"],
        ["reactvitemaket/src/app/data", "типы, mock-данные, адаптеры и store пользовательского состояния"],
        ["reactvitemaket/src/app/services", "API-клиент, поиск, подбор учреждений, напоминания"],
        ["src/backend/api", "маршруты FastAPI для public, user, admin, auth и статей"],
        ["src/backend/migrations", "SQL-миграции структуры SQLite базы"],
        ["src/backend/seeds", "начальные данные сценариев и справочного контента"],
        ["docs", "документация проекта, планы, отчёты, материалы для диплома"],
    ], [5.0, 10.8])


def build_doc():
    ensure_dirs()
    make_diagrams()
    prepare_screenshots()
    page_map = DEFAULT_TOC
    if PAGE_MAP_PATH.exists():
        try:
            page_map = {**DEFAULT_TOC, **json.loads(PAGE_MAP_PATH.read_text(encoding="utf-8"))}
        except Exception:
            page_map = DEFAULT_TOC
    doc = new_doc()
    title_page(doc)
    assignment_page(doc)
    abstract_page(doc)
    toc_page(doc, page_map)
    sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(2)
    set_page_start(sec, 5)
    add_page_number(sec)
    abbreviations(doc)
    introduction(doc)
    analytical_section(doc)
    practical_section(doc)
    user_guide(doc)
    economics(doc)
    labor(doc)
    conclusion_sources_appendix(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
